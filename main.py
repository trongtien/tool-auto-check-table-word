import re
import docx
from numpy import random, array as numpy_array
import pandas as pd
from docxtpl import DocxTemplate

'''
    index   0
    a       1
    b       2
    c       3
    d       4
    e       5
    f       6
    g       7
    h       8

'''

'''
1d      -> (1, 4)   : max -> 4 
2e      -> (2, 5)   : max -> 5
3c      -> (3, 3)   : max -> 3
4b      -> (4, 2)   : max -> 2
5a      -> (5, 1)   : max -> 2
6a      -> (6, 1)   : max -> 2
7e      -> (7, 5)   : max -> 5
8c      -> (8, 3)   : max -> 4
9d      -> (9, 4)   : max -> 4
10h     -> (10, 8)  : max -> 8
11e     -> (11, 5)  : max -> 5
12e     -> (12, 5)  : max -> 5
13d     -> (13, 4)  : max -> 4
14c     -> (12, 3)  : max -> 3
15a     -> (15, 1)  : max -> 2
16c     -> (16, 3)  : max -> 3
17e     -> (17, 5)  : max -> 5
18c     -> (18, 3)  : max -> 3
19d     -> (19, 4)  : max -> 4
20b     -> (20, 2)  : max -> 4
21c     -> (21, 3)  : max -> 3
22d     -> (22, 4)  : max -> 4
'''

enum_answer = {
    1: dict(selected_answer=4, max=4, cell=1),
    2: dict(selected_answer=5, max=5, cell=2),
    3: dict(selected_answer=3, max=3, cell=3),
    4: dict(selected_answer=2, max=2, cell=4),
    5: dict(selected_answer=1, max=2, cell=5),
    6: dict(selected_answer=1, max=2, cell=6),
    7: dict(selected_answer=5, max=5, cell=7),
    8: dict(selected_answer=3, max=4, cell=8),
    9: dict(selected_answer=4, max=4, cell=9),
    10: dict(selected_answer=8, max=8, cell=10),
    11: dict(selected_answer=5, max=5, cell=1),
    12: dict(selected_answer=5, max=5, cell=2),
    13: dict(selected_answer=4, max=4, cell=3),
    14: dict(selected_answer=3, max=3, cell=4),
    15: dict(selected_answer=1, max=2, cell=5),
    16: dict(selected_answer=3, max=3, cell=6),
    17: dict(selected_answer=5, max=5, cell=7),
    18: dict(selected_answer=3, max=3, cell=8),
    19: dict(selected_answer=4, max=4, cell=9),
    20: dict(selected_answer=2, max=4, cell=10),
    21: dict(selected_answer=3, max=3, cell=1),
    22: dict(selected_answer=4, max=4, cell=2)
}


def automationRandomSelectedAnswer(num_max, selected_answer):
    # 1-> num_max
    num_list_random = list(range(1, num_max + 1))
    num_random = random.choice(num_list_random)

    if num_random == selected_answer:
        return automationRandomSelectedAnswer(num_max, selected_answer)

    return num_random


def random_question_correl_answer(num_cell_root):
    question = []
    # 12 -> 18
    num_correl_answer_list_random = list(range(12, 19))
    num_correl_answer = random.choice(num_correl_answer_list_random)
    num_cell = num_cell_root


    for _ in range(num_correl_answer):
        num_cell_random = random.choice(num_cell)
        num_cell.remove(num_cell_random)

        question.append(num_cell_random)

    return question 

def no_accent_vietnamese(s: str) -> str:
    s = re.sub('[áàảãạăắằẳẵặâấầẩẫậ]', 'a', s)
    s = re.sub('[ÁÀẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬ]', 'A', s)
    s = re.sub('[éèẻẽẹêếềểễệ]', 'e', s)
    s = re.sub('[ÉÈẺẼẸÊẾỀỂỄỆ]', 'E', s)
    s = re.sub('[óòỏõọôốồổỗộơớờởỡợ]', 'o', s)
    s = re.sub('[ÓÒỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢ]', 'O', s)
    s = re.sub('[íìỉĩị]', 'i', s)
    s = re.sub('[ÍÌỈĨỊ]', 'I', s)
    s = re.sub('[úùủũụưứừửữự]', 'u', s)
    s = re.sub('[ÚÙỦŨỤƯỨỪỬỮỰ]', 'U', s)
    s = re.sub('[ýỳỷỹỵ]', 'y', s)
    s = re.sub('[ÝỲỶỸỴ]', 'Y', s)
    s = re.sub('đ', 'd', s)
    s = re.sub('Đ', 'D', s)
    return s

def gen_code(text: str) -> str:
        text = no_accent_vietnamese(text)

        while " " in text:
            text = text.replace(" ", "_")

        return text

def automationTick(index: str, path_doc_save_file: str, employee_import):
    print('==============================>')
    print('+++++++++++++++++++++++++++++')
    
    path = 'C:/Users/base/Documents/word/N4_template.docx'
    doc = docx.Document(path)
    name_path = gen_code(employee_import)
    print('[Name convert path]', name_path)
    file_name_suffix = name_path+"_BT-N4-"+index+".docx"
    value_default_cell='x'

    # Max row table tick
    question = random_question_correl_answer(list(range(1, 23)))
    print('Question tick', len(question))

    # paragraphs = doc.paragraphs

  
    # Get table row
    table = doc.tables[0]

    for numCell in list(range(1, 23)):
        selected_answer = enum_answer.get(numCell)['selected_answer']
        num_cell_set_value = enum_answer.get(numCell)['cell']


        max_answer = enum_answer.get(numCell)['max']
        selected_answer_random = automationRandomSelectedAnswer(max_answer, selected_answer)

        '''
            [11, 20] -> +11
            [21, 30] -> +21
        '''
        if numCell in list(range(11, 21)):
            selected_answer = selected_answer + 11
            selected_answer_random = selected_answer_random + 11

        if numCell in list(range(21, 31)):
            selected_answer = selected_answer + 21
            selected_answer_random = selected_answer_random + 21

        '''
            Add value table cell
        '''
        if numCell in question:
            print('(', num_cell_set_value,',',selected_answer, ')')
            table.cell(num_cell_set_value, selected_answer).text = value_default_cell
        else:
            print('(', num_cell_set_value,',',selected_answer_random, ')')
            table.cell(num_cell_set_value, selected_answer_random).text = value_default_cell

    
    doc.save(path_doc_save_file + file_name_suffix)

    question = []
    print('+++++++++++++++++++++++++++++')
    print("Save file success")
    print('==============================>')
    

def readExcel():
    path_read_template_excel = "C:/Users/base/Documents/word/MTT_DUKE_2024_DOT_2.xlsx"
    data_frame = pd.read_excel(path_read_template_excel)
    data_frame = pd.read_excel(path_read_template_excel, sheet_name="nhom_4_6")
    value_convert = data_frame.values

    arr_numpy = numpy_array(value_convert)
    arr_numpy[0:1]
    return arr_numpy

def main():
    path_save_file = 'C:/Users/base/Documents/word/convert/dot_2/nhom_4_6/'
    employees = readExcel()

    for idx, empl in enumerate(employees):
        employee_import = dict(USERNAME=empl[0], DATE=empl[2], GENDER=empl[1], NATIONAL=empl[3], CCCD=empl[5], POSITION=empl[6], JOB=empl[6], UNIT=empl[7])

        docx_template = DocxTemplate('C:/Users/base/Documents/word/BT-N4.docx')
        print(f"========= >{employee_import['USERNAME']} | {employee_import['DATE']} | {employee_import['GENDER']} | {employee_import['NATIONAL']} | {employee_import['CCCD']} | {employee_import['POSITION']} | {employee_import['USERNAME']} | {employee_import['UNIT']} < ============")
        context = {
                'USERNAME': employee_import['USERNAME'],
                # 'DATE':  f"{'{0:.0f}'.format(employee_import['DATE'])}    ",
                'DATE':  f"{employee_import['DATE']}    ",
                'GENDER': employee_import['GENDER'],
                'NATIONAL': employee_import['NATIONAL'],
                # 'CCCD': '{0:.0f}'.format(employee_import['CCCD']),
                'CCCD': employee_import['CCCD'],
                'POSITION': employee_import['POSITION'],
                'UNIT': f"{employee_import['UNIT']}",
            }     
        
        docx_template.render(context)
        docx_template.save('C:/Users/base/Documents/word/N4_template.docx')
        automationTick(str(idx+1), path_save_file, empl[0])



    # num_word_generator = 1
    # path_save_file = 'C:/Users/base/Documents/word/convert/'
    # for index in range(num_word_generator):
    #     automationTick(str(index+1), path_save_file, employees)

main()